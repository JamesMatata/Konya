import json
import logging
import os
import re

import requests
from bs4 import BeautifulSoup
from django.utils.translation import gettext as _
from django.utils.translation import override

from accounts.i18n import LANGUAGE_DISPLAY_NAMES, normalize_language_code

logger = logging.getLogger(__name__)

FEATHERLESS_API_URL = "https://api.featherless.ai/v1/chat/completions"
DEFAULT_MODEL = "meta-llama/Meta-Llama-3.1-8B-Instruct"
MAX_POLICY_TEXT_LENGTH = 12000
MAX_URL_TEXT_LENGTH = 8000
MAX_TITLE_LENGTH = 60


def get_ai_error_message(language_code="en"):
    with override(normalize_language_code(language_code)):
        return _("I'm sorry, I lost my connection for a moment. Could we try that again?")


def language_name_for_prompt(language_code):
    """Human-readable language label for LLM system/user directives."""
    code = normalize_language_code(language_code)
    if code == "sheng":
        return "Sheng (informal Kenyan urban language mixing English and Kiswahili)"
    return LANGUAGE_DISPLAY_NAMES.get(code, "English")


# Backwards-compatible alias for internal references during migration.
AI_ERROR_MESSAGE = get_ai_error_message("en")

TITLE_SYSTEM_PROMPT = """You generate short conversation titles for a social policy navigation app.
Read the user's opening message and return a concise title of 3 to 5 words that summarizes their topic.
Examples: "Maternity Leave Eligibility", "Nairobi Unemployment Support", "NHIF Registration Help"

Respond with valid JSON only — no markdown fences:
{"title": "Your Title Here"}"""

GATEKEEPER_REJECTION_PREFIXES = (
    "The document you provided does not seem",
    "This document also doesn't seem to match",
    "We've tried several documents",
    "I'm sorry, I lost my connection",
)


def is_gatekeeper_message(message):
    """True for Phase 1 rejection/error messages that must not feed Phase 2."""
    if getattr(message, "is_gatekeeper", False):
        return True
    if message.role != "ai":
        return False
    if message.is_error:
        return True
    return any(message.content.startswith(prefix) for prefix in GATEKEEPER_REJECTION_PREFIXES)


def _is_gatekeeper_message(message):
    return is_gatekeeper_message(message)


def _count_prior_navigator_ai_messages(chat_instance):
    """Count Phase 2 navigator responses only (excludes gatekeeper rejections)."""
    count = 0
    for message in chat_instance.messages.filter(role="ai", is_error=False).order_by("timestamp"):
        if not _is_gatekeeper_message(message):
            count += 1
    return count


def _truncate(text, limit):
    if len(text) <= limit:
        return text
    return text[:limit].rstrip() + "\n\n[Content truncated for length.]"


def _fetch_url_text(url):
    try:
        response = requests.get(
            url,
            timeout=20,
            headers={"User-Agent": "Konya/1.0 (+https://localhost)"},
        )
        response.raise_for_status()
    except requests.RequestException as exc:
        logger.warning("Failed to fetch policy URL %s: %s", url, exc)
        return f"[Could not retrieve content from {url}.]"

    content_type = response.headers.get("Content-Type", "")
    if "html" in content_type or response.text.lstrip().startswith("<"):
        soup = BeautifulSoup(response.text, "html.parser")
        for tag in soup(["script", "style", "noscript"]):
            tag.decompose()
        text = soup.get_text(separator="\n")
        text = re.sub(r"\n{3,}", "\n\n", text)
        text = re.sub(r"[ \t]+", " ", text).strip()
        return _truncate(text, MAX_URL_TEXT_LENGTH)

    return _truncate(response.text.strip(), MAX_URL_TEXT_LENGTH)


def _read_text_upload(uploaded_file):
    uploaded_file.seek(0)
    raw = uploaded_file.read()
    uploaded_file.seek(0)
    if isinstance(raw, bytes):
        return raw.decode("utf-8", errors="replace")
    return str(raw)


def _read_docx_upload(uploaded_file):
    from docx import Document

    uploaded_file.seek(0)
    document = Document(uploaded_file)
    uploaded_file.seek(0)
    paragraphs = [paragraph.text.strip() for paragraph in document.paragraphs if paragraph.text.strip()]
    return "\n".join(paragraphs).strip()


def _extract_file_text(message):
    if not message.attached_file:
        return ""

    filename = message.attached_file.name
    extension = filename.rsplit(".", 1)[-1].lower() if "." in filename else ""

    try:
        if extension == "pdf":
            from pypdf import PdfReader

            reader = PdfReader(message.attached_file.path)
            pages = [page.extract_text() or "" for page in reader.pages]
            text = "\n".join(pages).strip()
            if text:
                return _truncate(text, MAX_POLICY_TEXT_LENGTH)
            return f"[PDF uploaded ({filename}) but no extractable text was found.]"

        if extension == "txt":
            with message.attached_file.open("rb") as handle:
                text = handle.read().decode("utf-8", errors="replace").strip()
            if text:
                return _truncate(text, MAX_POLICY_TEXT_LENGTH)
            return f"[Text file uploaded ({filename}) but it appears to be empty.]"

        if extension == "docx":
            from docx import Document

            document = Document(message.attached_file.path)
            paragraphs = [paragraph.text.strip() for paragraph in document.paragraphs if paragraph.text.strip()]
            text = "\n".join(paragraphs).strip()
            if text:
                return _truncate(text, MAX_POLICY_TEXT_LENGTH)
            return f"[Word document uploaded ({filename}) but no extractable text was found.]"

        if extension == "doc":
            return (
                f"[Legacy Word file uploaded: {filename}. Please resubmit as .docx or PDF if possible.]"
            )

        if extension in {"jpg", "jpeg", "png", "gif", "webp", "heic"}:
            return (
                f"[Image uploaded: {filename}. Visual policy text was not automatically extracted; "
                "use the user's written context and ask only for policy-required facts not yet provided.]"
            )
    except Exception as exc:
        logger.warning("Failed to extract file text from %s: %s", filename, exc)
        return f"[Could not read uploaded file: {filename}.]"

    return f"[Unsupported uploaded file type: {filename}.]"


def _extract_uploaded_file_text(uploaded_file):
    """Extract text from an in-memory uploaded file before it is saved to storage."""
    if not uploaded_file:
        return ""

    filename = uploaded_file.name
    extension = filename.rsplit(".", 1)[-1].lower() if "." in filename else ""

    try:
        if extension == "pdf":
            from pypdf import PdfReader

            uploaded_file.seek(0)
            reader = PdfReader(uploaded_file)
            pages = [page.extract_text() or "" for page in reader.pages]
            text = "\n".join(pages).strip()
            uploaded_file.seek(0)
            if text:
                return _truncate(text, MAX_POLICY_TEXT_LENGTH)
            return f"[PDF uploaded ({filename}) but no extractable text was found.]"

        if extension == "txt":
            text = _read_text_upload(uploaded_file).strip()
            if text:
                return _truncate(text, MAX_POLICY_TEXT_LENGTH)
            return f"[Text file uploaded ({filename}) but it appears to be empty.]"

        if extension == "docx":
            text = _read_docx_upload(uploaded_file).strip()
            if text:
                return _truncate(text, MAX_POLICY_TEXT_LENGTH)
            return f"[Word document uploaded ({filename}) but no extractable text was found.]"

        if extension == "doc":
            return (
                f"[Legacy Word file uploaded: {filename}. Please resubmit as .docx or PDF if possible.]"
            )

        if extension in {"jpg", "jpeg", "png", "gif", "webp", "heic"}:
            return (
                f"[Image uploaded: {filename}. Visual policy text was not automatically extracted; "
                "use the user's written context and ask only for policy-required facts not yet provided.]"
            )
    except Exception as exc:
        logger.warning("Failed to extract uploaded file text from %s: %s", filename, exc)
        return f"[Could not read uploaded file: {filename}.]"

    return f"[Unsupported uploaded file type: {filename}.]"


def _build_policy_preview(attached_url=None, attached_file=None):
    sections = []
    if attached_url:
        url_text = _fetch_url_text(attached_url)
        sections.append(f"Reference link ({attached_url}):\n{url_text}")
    if attached_file:
        file_text = _extract_uploaded_file_text(attached_file)
        sections.append(f"Uploaded document ({attached_file.name}):\n{file_text}")
    return "\n\n".join(sections)


GATEKEEPER_SYSTEM_PROMPT = """You are a strict relevance evaluator. You will be provided with a User's Situation and the text of a Policy Document/URL.

Your ONLY job is to determine whether the document is RELEVANT to the user's situation — not whether the user qualifies, wins, or gets a favorable outcome.

=== THE RELEVANCE DEFINITION RULE ===

A document is strictly RELEVANT (is_relevant: true) if it contains the rules, criteria, eligibility conditions, restrictions, or facts needed to answer the user's situation — even partially.

CRITICAL: A document is still 100% relevant even if those rules disqualify the user or answer their question with a "no".

Examples of RELEVANT (is_relevant: true):
- User asks "Can I compete alone?" and the document states "Solo participants are not allowed" → RELEVANT (definitive answer).
- User asks about maternity leave and the document lists who is NOT eligible → RELEVANT (eligibility rules apply).
- User wants a grant and the document explains the program requirements, even if the user likely fails them → RELEVANT.

Examples of NOT RELEVANT (is_relevant: false):
- The document is about a completely different program, country, or topic than the user's situation.
- The document is generic marketing with no rules, criteria, or factual guidance.
- No readable policy text could be matched to the user's stated need.

Do NOT reject a document because:
- The user would be ineligible or disqualified.
- The answer to their question is negative or unfavorable.
- The rules contradict what the user hopes to do.

Reject ONLY when the document does not contain the kind of rules or facts needed to address their situation at all.

Return a strict JSON object: {"is_relevant": true/false, "reason": "A one-sentence explanation"}

In "reason", describe topical fit only (e.g. "contains team-participation rules that directly answer the solo-entry question") — do not judge whether the user qualifies.

Respond with valid JSON only — no markdown fences, no prose outside the JSON object."""


def _parse_gatekeeper_response(raw_text, *, language_code="en"):
    language_code = normalize_language_code(language_code)
    with override(language_code):
        parse_error_reason = _("Could not parse the relevance evaluation.")

    text = raw_text.strip()
    try:
        data = json.loads(text)
    except json.JSONDecodeError:
        brace_match = re.search(r"(\{.*\})", text, re.DOTALL)
        if brace_match:
            try:
                data = json.loads(brace_match.group(1))
            except json.JSONDecodeError:
                return {"is_relevant": False, "reason": parse_error_reason}
        else:
            return {"is_relevant": False, "reason": parse_error_reason}

    is_relevant = data.get("is_relevant")
    if is_relevant is None:
        is_relevant = data.get("valid")

    return {
        "is_relevant": bool(is_relevant),
        "reason": str(data.get("reason", "")).strip(),
    }


def build_gatekeeper_rejection_message(invalid_doc_attempts, reason="", language_code="en"):
    """Build the user-facing rejection copy after a failed Gatekeeper check."""
    language_code = normalize_language_code(language_code)
    with override(language_code):
        if invalid_doc_attempts == 1:
            reason_text = reason.strip().rstrip(".") or _(
                "it does not appear to address your stated situation"
            )
            return _(
                "The document you provided does not seem to contain the right information for your "
                "situation because %(reason)s."
            ) % {"reason": reason_text}

        if invalid_doc_attempts == 2:
            return _(
                "This document also doesn't seem to match. Try searching the official portal for your "
                "program or looking for eligibility guidelines and application requirements."
            )

        return _(
            "We've tried several documents that still don't match your situation. Please look for the "
            "official program page, eligibility guidelines, or application requirements from a trusted "
            "government or institution source."
        )


def build_empty_planner_message(language_code="en"):
    with override(normalize_language_code(language_code)):
        return _(
            "I couldn't extract clear eligibility rules from this document. Please try a more complete "
            "policy page, official program guidelines, or a readable PDF/Word document."
        )


def evaluate_policy_document(chat_instance, *, attached_url=None, attached_file=None, language_code="en"):
    """
    Phase 1 — Gatekeeper: lightweight relevance check only.

    Does NOT generate conversational navigation responses.
    Returns dict: is_relevant (bool), reason (str), error (bool).
    """
    language_code = normalize_language_code(language_code)
    language_name = language_name_for_prompt(language_code)

    first_user_message = chat_instance.messages.filter(role="user").order_by("timestamp").first()
    situation = first_user_message.content.strip() if first_user_message else ""
    policy_preview = _build_policy_preview(attached_url=attached_url, attached_file=attached_file)

    if not policy_preview.strip():
        with override(language_code):
            return {
                "is_relevant": False,
                "reason": _("no readable policy text could be extracted from the submission"),
                "error": False,
            }

    user_prompt = (
        f"Language directive: write the \"reason\" field entirely in {language_name}.\n\n"
        "Evaluate TOPICAL RELEVANCE only. Mark is_relevant: true if the policy text contains rules, "
        "criteria, restrictions, or facts that could answer the user's situation — including when those "
        "rules say 'no' or disqualify the user. Mark is_relevant: false only if the document is the wrong "
        "topic or lacks usable rules for their situation.\n\n"
        f"=== USER SITUATION ===\n{situation}\n\n"
        f"=== POLICY DOCUMENT / URL TEXT ===\n{policy_preview}"
    )

    messages = [
        {"role": "system", "content": GATEKEEPER_SYSTEM_PROMPT},
        {"role": "user", "content": user_prompt},
    ]

    try:
        raw_content = _call_featherless(messages, max_tokens=256, temperature=0.1)
        parsed = _parse_gatekeeper_response(raw_content, language_code=language_code)
        return {
            "is_relevant": parsed["is_relevant"],
            "reason": parsed["reason"],
            "error": False,
        }
    except (ValueError, RuntimeError):
        logger.exception("Gatekeeper evaluation failed for chat %s", chat_instance.id)
        return {
            "is_relevant": False,
            "reason": "",
            "error": True,
        }


PLANNER_SYSTEM_PROMPT = """You are a policy analyst. Read the provided Policy Document. Extract ONLY the logical \
conditions required to determine if a person is eligible for this program. Do NOT extract application logistics \
(like requiring an ID, bank account, or signature)—those are for later.

You will also receive the user's initial situation. For each qualifying rule return:
- key: snake_case identifier
- label: plain-language question or requirement (for the user)
- policy_source: short quote or paraphrase of the policy line
- value: null if unknown from the situation, true if clearly met, false if clearly not met

Return strict JSON only — no markdown fences:
{
  "criteria": [
    {
      "key": "age_requirement_met",
      "label": "Whether you are between 18 and 65 years old",
      "policy_source": "Applicants must be aged 18–65",
      "value": null
    }
  ]
}"""


def planner_system_prompt(language_code="en"):
    language_name = language_name_for_prompt(language_code)
    return (
        PLANNER_SYSTEM_PROMPT
        + f"\n\nLANGUAGE DIRECTIVE: Write every criterion label in {language_name}. "
        "Use natural, conversational wording. Keep keys as snake_case English identifiers only."
    )


EXTRACTOR_SYSTEM_PROMPT = """You extract eligibility answers from the user's conversation (latest message first, then earlier context).

You receive a checklist of criteria (some still unknown) with labels and policy context.
You may also receive the last question the assistant asked — treat the user's message as an answer to that question when it fits.
Return strict JSON only — no markdown fences:
{"updates": {"criterion_key": true, "another_key": false}}

Rules:
- Only include keys you can determine from the user's message with reasonable confidence.
- Omit keys that remain unknown — do not include them in updates.
- Use true if the user meets the criterion, false if they clearly do not.
- Do not invent facts the user did not state.
- The user may answer several criteria in one message — extract every key you can from that message.
- Treat clear intent as an answer even when informal or partial:
  - "yes", "yeah", "yep", "I do", "correct" → true for the criterion being asked.
  - "no", "nope", "not really", "I don't" → false for the criterion being asked.
  - "I am 22", "22 years", "22 years old", "im 22" → resolve age criteria using the threshold in the label/policy (e.g. at least 14 → true).
  - "I can ask a friend to join my team" → true for team-participation criteria.
  - "I am a student graduating in October" → answers enrollment/education criteria.
- Numbers in the user's message often answer age, income, household size, or duration criteria — compare them to the requirement in the label.
- If the user gives a number or short reply right after a specific question, map it to that criterion first.
- You may receive the FULL conversation so far — extract from the latest message first, then use earlier messages for anything still pending.
- Treat negations and informal phrasing as real answers:
  - "I don't have a team", "no team", "solo for now" → false for team-participation / team-size criteria.
  - "I am a student at [school/university]" → true for enrollment/education criteria.
  - "not employed", "unemployed", "I'm a student" → true for "not employed full-time" style criteria (they are not in that employment).
- If team participation is clearly false, set team-member / all-members-team criteria to false when they only apply to teams.
- Map partial answers across every pending criterion the user may be addressing — do not wait for one criterion at a time."""


def _sanitize_checklist_value(value):
    if value is None:
        return None
    if isinstance(value, bool):
        return value
    if isinstance(value, str):
        normalized = value.lower().strip()
        if normalized in {"true", "yes", "met", "eligible"}:
            return True
        if normalized in {"false", "no", "not met", "ineligible"}:
            return False
        if normalized in {"null", "unknown", "none", "unclear", ""}:
            return None
    return None


def _sanitize_checklist(data):
    if not isinstance(data, dict):
        return {}

    payload = data.get("checklist") if isinstance(data.get("checklist"), dict) else data
    checklist = {}
    for key, value in payload.items():
        if not isinstance(key, str):
            continue
        safe_key = re.sub(r"[^a-z0-9_]+", "_", key.lower().strip()).strip("_")
        if not safe_key:
            continue
        checklist[safe_key] = _sanitize_checklist_value(value)
    return checklist


def _safe_criterion_key(key):
    if not isinstance(key, str):
        return ""
    return re.sub(r"[^a-z0-9_]+", "_", key.lower().strip()).strip("_")


def _parse_planner_response(raw_text):
    text = raw_text.strip()
    data = None
    try:
        data = json.loads(text)
    except json.JSONDecodeError:
        fence_match = re.search(r"```(?:json)?\s*(\{.*?\})\s*```", text, re.DOTALL)
        if fence_match:
            try:
                data = json.loads(fence_match.group(1))
            except json.JSONDecodeError:
                data = None
        if data is None:
            brace_match = re.search(r"(\{.*\})", text, re.DOTALL)
            if brace_match:
                try:
                    data = json.loads(brace_match.group(1))
                except json.JSONDecodeError:
                    data = None

    if not isinstance(data, dict):
        return {}, {}

    checklist = {}
    criteria_meta = {}

    if isinstance(data.get("criteria"), list):
        for item in data["criteria"]:
            if not isinstance(item, dict):
                continue
            key = _safe_criterion_key(item.get("key", ""))
            if not key:
                continue
            checklist[key] = _sanitize_checklist_value(item.get("value"))
            criteria_meta[key] = {
                "label": str(item.get("label", "")).strip() or key.replace("_", " ").capitalize(),
                "policy_source": str(item.get("policy_source", "")).strip(),
            }
        return checklist, criteria_meta

    flat = _sanitize_checklist(data)
    for key in flat:
        criteria_meta[key] = {"label": key.replace("_", " ").capitalize(), "policy_source": ""}
    return flat, criteria_meta


def _get_criterion_label(key, criteria_meta=None):
    meta = (criteria_meta or {}).get(key) or {}
    label = meta.get("label", "").strip()
    if label:
        return label
    return key.replace("_", " ").strip().capitalize()


def _criterion_search_text(key, criteria_meta=None):
    meta = (criteria_meta or {}).get(key) or {}
    label = meta.get("label", "")
    policy = meta.get("policy_source", "")
    return f"{key} {label} {policy}".lower()


def _get_thanks_label(key, criteria_meta=None):
    """Short human label for acknowledgement copy — not the full policy question."""
    label = _get_criterion_label(key, criteria_meta).strip().rstrip("?")
    lowered = label.lower()
    prefixes = (
        "whether you are ",
        "whether you ",
        "whether ",
        "are you ",
        "will you ",
        "will all ",
        "do you ",
        "have you ",
    )
    for prefix in prefixes:
        if lowered.startswith(prefix):
            label = label[len(prefix) :].strip()
            lowered = label.lower()
    if label:
        return label[0].upper() + label[1:]
    return key.replace("_", " ").strip().capitalize()


def _is_team_membership_question(key, criteria_meta=None):
    text = _criterion_search_text(key, criteria_meta)
    return any(
        phrase in text
        for phrase in (
            "team member",
            "members of your team",
            "members of the team",
            "all team members",
            "each team member",
            "every team member",
            "all members meet",
            "team members meet",
        )
    )


def _is_team_participation_question(key, criteria_meta=None):
    if _is_team_membership_question(key, criteria_meta):
        return False
    text = _criterion_search_text(key, criteria_meta)
    if "team" not in text:
        return False
    return any(
        phrase in text
        for phrase in (
            "participat",
            "team of",
            "2-5",
            "2–5",
            "forming a team",
            "as a team",
        )
    )


def _team_participation_is_false(checklist, criteria_meta=None):
    for key, value in (checklist or {}).items():
        if value is not False:
            continue
        if _is_team_participation_question(key, criteria_meta):
            return True
        text = _criterion_search_text(key, criteria_meta)
        if (
            "team" in text
            and not _is_team_membership_question(key, criteria_meta)
            and any(word in text for word in ("participat", "team of", "2-5", "2–5"))
        ):
            return True
    return False


def _get_askable_null_keys(checklist, criteria_meta=None):
    """Null criteria that should still be asked — excludes dependents made irrelevant."""
    null_keys = _null_checklist_keys(checklist)
    if not null_keys:
        return []
    if not _team_participation_is_false(checklist, criteria_meta):
        return null_keys
    return [key for key in null_keys if not _is_team_membership_question(key, criteria_meta)]


def _is_negated_eligibility_question(key, criteria_meta=None):
    """Criteria phrased as 'not employed' / 'not disqualified' — avoid naive yes/no mapping."""
    text = _criterion_search_text(key, criteria_meta).replace("_", " ")
    return bool(
        re.search(
            r"\bnot\b.{0,40}\b(employ|working|work\s+full|disqualif|ineligible|restricted)",
            text,
        )
    )


def _apply_derived_checklist_resolutions(checklist, criteria_meta=None, resolved_sources=None):
    """Fill dependent criteria made irrelevant by parent answers (e.g. no team → members N/A)."""
    checklist = dict(checklist or {})
    resolved_sources = dict(resolved_sources or {})
    derived = _infer_dependent_criteria_updates(checklist, criteria_meta)
    for key, value in derived.items():
        if checklist.get(key) is None:
            checklist[key] = value
            resolved_sources[key] = "derived"
    return checklist, resolved_sources


def _merge_extractor_updates(heuristic_updates, llm_updates):
    """Heuristic extractions win over LLM when both supply a value for the same key."""
    merged = dict(llm_updates or {})
    merged.update(heuristic_updates or {})
    return merged


def _interview_needs_more_questions(state):
    """True while askable unknown criteria remain after dependency resolution."""
    checklist = dict((state or {}).get("checklist") or {})
    criteria_meta = (state or {}).get("criteria_meta") or {}
    resolved_sources = dict((state or {}).get("resolved_sources") or {})
    checklist, _ = _apply_derived_checklist_resolutions(
        checklist, criteria_meta, resolved_sources
    )
    return bool(_get_askable_null_keys(checklist, criteria_meta))


def _infer_dependent_criteria_updates(checklist, criteria_meta=None):
    """Derive checklist updates when parent answers make child criteria irrelevant."""
    updates = {}
    if not _team_participation_is_false(checklist, criteria_meta):
        return updates
    for key in _null_checklist_keys(checklist):
        if _is_team_membership_question(key, criteria_meta):
            updates[key] = False
    return updates


def _gather_user_conversation_text(chat_instance):
    parts = []
    for message in chat_instance.messages.filter(role="user").order_by("timestamp"):
        text = (message.content or "").strip()
        if text:
            parts.append(text)
    return "\n\n".join(parts)


def _opening_situation_text(chat_instance):
    first_user_message = chat_instance.messages.filter(role="user").order_by("timestamp").first()
    return first_user_message.content.strip() if first_user_message else ""


def _checklist_to_navigator_fields(checklist, criteria_meta=None, resolved_sources=None):
    """Derive navigator criteria/missing_items from the planner checklist."""
    criteria_meta = criteria_meta or {}
    resolved_sources = resolved_sources or {}
    criteria = []
    missing_items = []
    askable_null_keys = _get_askable_null_keys(checklist, criteria_meta)

    for key, value in checklist.items():
        label = _get_criterion_label(key, criteria_meta)
        meta = criteria_meta.get(key, {})
        if value is None:
            if key in askable_null_keys:
                status = "missing"
                value_text = ""
                missing_items.append(key)
            else:
                status = "skipped"
                value_text = ""
        elif value is True:
            status = "known"
            value_text = "met"
        else:
            status = "disqualified"
            value_text = "not met"

        criteria.append(
            {
                "key": key,
                "label": label,
                "status": status,
                "value": value_text,
                "policy_source": meta.get("policy_source", ""),
            }
        )

    unknown_count = len(missing_items)
    if unknown_count:
        summary = f"{unknown_count} qualifying rule(s) still need verification."
    else:
        summary = "All qualifying rules resolved from policy and initial situation."

    return {
        "planner_completed": True,
        "checklist": checklist,
        "criteria_meta": criteria_meta,
        "resolved_sources": dict(resolved_sources),
        "conclusion_reached": False,
        "phase": "clarification",
        "bulk_intake_completed": False,
        "summary": summary,
        "candidate_programs": [],
        "criteria": criteria,
        "missing_items": missing_items,
        "questions_asked": [],
        "pending_ask_keys": [],
        "last_asked_key": "",
        "next_target_variable": askable_null_keys[0] if askable_null_keys else "",
        "action_blueprint": "",
    }


def _build_state_from_checklist(existing_state, checklist, *, criteria_meta=None, resolved_sources=None):
    existing = existing_state if isinstance(existing_state, dict) else {}
    meta = criteria_meta if criteria_meta is not None else existing.get("criteria_meta") or {}
    sources = resolved_sources if resolved_sources is not None else existing.get("resolved_sources") or {}

    new_state = _checklist_to_navigator_fields(checklist, meta, sources)
    for field in (
        "planner_completed",
        "interviewer_bulk_completed",
        "navigation_complete",
        "questions_asked",
        "pending_ask_keys",
        "last_asked_key",
    ):
        if field in existing:
            new_state[field] = existing[field]
    new_state["questions_asked"] = list(existing.get("questions_asked") or [])
    new_state["pending_ask_keys"] = list(existing.get("pending_ask_keys") or [])
    new_state["last_asked_key"] = existing.get("last_asked_key") or ""
    return _normalize_eligibility_state(new_state)


def finalize_eligibility_state(state):
    """Apply dependency resolution and rebuild navigator-facing fields."""
    existing = dict(state or {})
    checklist = dict(existing.get("checklist") or {})
    criteria_meta = existing.get("criteria_meta") or {}
    resolved_sources = dict(existing.get("resolved_sources") or {})
    checklist, resolved_sources = _apply_derived_checklist_resolutions(
        checklist, criteria_meta, resolved_sources
    )
    return _build_state_from_checklist(
        existing,
        checklist,
        criteria_meta=criteria_meta,
        resolved_sources=resolved_sources,
    )


def apply_checklist_updates(state, updates, *, source="conversation"):
    """Code-owned merge of extractor updates into eligibility_state."""
    existing = state if isinstance(state, dict) else {}
    checklist = dict(existing.get("checklist") or {})
    criteria_meta = dict(existing.get("criteria_meta") or {})
    resolved_sources = dict(existing.get("resolved_sources") or {})

    for key, value in (updates or {}).items():
        if key not in checklist:
            continue
        sanitized = _sanitize_checklist_value(value)
        if sanitized is None:
            continue
        previous = checklist.get(key)
        checklist[key] = sanitized
        if previous is None:
            resolved_sources[key] = source

    checklist, resolved_sources = _apply_derived_checklist_resolutions(
        checklist, criteria_meta, resolved_sources
    )

    return _build_state_from_checklist(
        existing,
        checklist,
        criteria_meta=criteria_meta,
        resolved_sources=resolved_sources,
    )


def is_ready_for_blueprint(state):
    """All criteria resolved with a tracked source — safe to generate Action Blueprint."""
    if _checklist_has_nulls(state):
        return False
    checklist = (state or {}).get("checklist") or {}
    if not checklist:
        return False
    resolved_sources = (state or {}).get("resolved_sources") or {}
    return all(key in resolved_sources for key in checklist)


def cache_policy_text_for_chat(chat_instance, *, attached_url=None, attached_file=None):
    """Persist extracted policy text once after Gatekeeper approval."""
    preview = _build_policy_preview(attached_url=attached_url, attached_file=attached_file)
    if not preview.strip():
        first_user = chat_instance.messages.filter(role="user").order_by("timestamp").first()
        if first_user:
            preview = _build_verified_policy_document_from_message(first_user)
    if preview.strip():
        chat_instance.cached_policy_text = preview
        chat_instance.save(update_fields=["cached_policy_text"])
    return preview


def _build_verified_policy_document_from_message(first_user_message):
    sections = []
    if first_user_message.attached_url:
        url_text = _fetch_url_text(first_user_message.attached_url)
        sections.append(f"Source URL: {first_user_message.attached_url}\n\n{url_text}")
    if first_user_message.attached_file:
        file_text = _extract_file_text(first_user_message)
        filename = first_user_message.attached_file.name.rsplit("/", 1)[-1]
        sections.append(f"Source file: {filename}\n\n{file_text}")
    return "\n\n---\n\n".join(sections)


def _parse_age_from_message(text):
    """Extract a plausible age in years from informal user text."""
    if not text:
        return None
    lowered = text.lower().strip()
    patterns = (
        r"(?:i am|i'm|im)\s*(\d{1,3})\b",
        r"\b(\d{1,3})\s*(?:years?\s*old|yrs?\b|years?\b)",
        r"\bage\s*(?:is\s*)?(\d{1,3})\b",
    )
    for pattern in patterns:
        match = re.search(pattern, lowered)
        if not match:
            continue
        age = int(match.group(1))
        if 0 < age < 130:
            return age
    return None


def _parse_age_rule_from_text(*parts):
    """Return ('min', n), ('max', n), or ('range', low, high) from criterion copy."""
    text = " ".join(str(part or "") for part in parts).lower().replace("_", " ")
    if not text.strip():
        return None

    match = re.search(r"at least (\d{1,3})", text)
    if match:
        return ("min", int(match.group(1)))

    match = re.search(r"(?:minimum|min\.?)\s*(?:age\s*)?(?:of\s*)?(\d{1,3})", text)
    if match:
        return ("min", int(match.group(1)))

    match = re.search(r"(\d{1,3})\s*(?:years?|yrs?)\s*(?:or older|and (?:over|above)|\+)", text)
    if match:
        return ("min", int(match.group(1)))

    match = re.search(r"between (\d{1,3})\s*(?:and|-|to)\s*(\d{1,3})", text)
    if match:
        low, high = int(match.group(1)), int(match.group(2))
        return ("range", min(low, high), max(low, high))

    match = re.search(r"(?:under|below|less than)\s*(\d{1,3})", text)
    if match:
        return ("max", int(match.group(1)))

    match = re.search(r"aged?\s*(\d{1,3})\s*[-–]\s*(\d{1,3})", text)
    if match:
        low, high = int(match.group(1)), int(match.group(2))
        return ("range", min(low, high), max(low, high))

    return None


def _evaluate_age_against_rule(age, rule):
    if rule[0] == "min":
        return age >= rule[1]
    if rule[0] == "max":
        return age < rule[1]
    if rule[0] == "range":
        return rule[1] <= age <= rule[2]
    return None


def _heuristic_checklist_updates(
    user_message,
    checklist,
    criteria_meta,
    *,
    focus_key="",
    pending_ask_keys=None,
    conversation_text="",
):
    """Fast path for obvious yes/no and age answers before the LLM extractor."""
    pending_ask_keys = list(pending_ask_keys or [])
    texts_to_scan = []
    if conversation_text and conversation_text.strip():
        texts_to_scan.append(conversation_text.strip())
    if user_message.strip():
        latest = user_message.strip()
        if latest not in texts_to_scan:
            texts_to_scan.append(latest)

    merged = {}
    working = dict(checklist or {})
    for text in texts_to_scan:
        partial = _heuristic_checklist_updates_from_text(
            text,
            working,
            criteria_meta,
            focus_key=focus_key,
            pending_ask_keys=pending_ask_keys,
        )
        merged.update(partial)
        working.update(partial)
    return merged


def _heuristic_checklist_updates_from_text(
    user_message,
    checklist,
    criteria_meta,
    *,
    focus_key="",
    pending_ask_keys=None,
):
    updates = {}
    null_keys = _null_checklist_keys(checklist)
    if not user_message.strip() or not null_keys:
        return updates

    pending_ask_keys = [key for key in (pending_ask_keys or []) if key in null_keys]
    text = user_message.strip()
    text_lower = text.lower()
    normalized = re.sub(r"[^\w\s']", "", text_lower).strip()

    short_answer_keys = pending_ask_keys or ([focus_key] if focus_key in null_keys else list(null_keys))
    if len(short_answer_keys) == 1:
        key = short_answer_keys[0]
        negated = _is_negated_eligibility_question(key, criteria_meta)
        affirm = {
            "yes",
            "yeah",
            "yep",
            "yup",
            "correct",
            "true",
            "affirmative",
            "i do",
            "definitely",
            "sure",
            "absolutely",
            "of course",
        }
        deny = {"no", "nope", "nah", "false", "negative", "i don't", "i dont", "not really", "not"}
        if not negated:
            if normalized in affirm or re.match(r"^yes\b", normalized):
                updates[key] = True
            elif normalized in deny or (
                re.match(r"^no\b", normalized) and len(normalized.split()) <= 4
            ):
                updates[key] = False

    age = _parse_age_from_message(text)
    if age is not None:
        candidate_keys = list(null_keys)
        if focus_key in null_keys:
            candidate_keys = [focus_key] + [k for k in null_keys if k != focus_key]
        for key in candidate_keys:
            if key in updates:
                continue
            meta = criteria_meta.get(key) or {}
            label = _get_criterion_label(key, criteria_meta)
            policy_source = meta.get("policy_source", "")
            combined = f"{key} {label} {policy_source}".lower()
            rule = _parse_age_rule_from_text(label, policy_source, key)
            if not rule and "age" not in combined:
                continue
            if not rule:
                continue
            evaluated = _evaluate_age_against_rule(age, rule)
            if evaluated is not None:
                updates[key] = evaluated

    if re.search(
        r"(?:don'?t|do not) have (?:a )?team|no team|without a team|solo(?: for now)?|on my own for now|for now i don'?t have",
        text_lower,
    ):
        for key in null_keys:
            if key in updates:
                continue
            if _is_team_participation_question(key, criteria_meta):
                updates[key] = False

    if re.search(
        r"\b(student|enrolled|enrolment|enrollment|university|college|school|multimedia)\b",
        text_lower,
    ):
        for key in null_keys:
            if key in updates:
                continue
            criterion_text = _criterion_search_text(key, criteria_meta)
            if any(
                word in criterion_text
                for word in (
                    "enrol",
                    "student",
                    "school",
                    "university",
                    "college",
                    "doctoral",
                    "graduate",
                    "accredited",
                    "secondary",
                )
            ):
                updates[key] = True

    if re.search(
        r"(not employed|unemployed|no(?:t)? employed|i'?m a student|i am a student|don'?t work|do not work|no i am not employed)",
        text_lower,
    ):
        for key in null_keys:
            if key in updates:
                continue
            criterion_text = _criterion_search_text(key, criteria_meta)
            if "employ" in criterion_text and "not" in criterion_text:
                updates[key] = True

    return updates


def extract_checklist_updates(chat_instance, user_message, state, language_code="en"):
    """Extractor agent — structured answer parsing only, no user-facing text."""
    language_code = normalize_language_code(language_code)
    checklist = (state or {}).get("checklist") or {}
    criteria_meta = (state or {}).get("criteria_meta") or {}
    null_keys = _null_checklist_keys(checklist)
    if not user_message.strip() or not null_keys:
        return {"updates": {}, "error": False}

    pending = []
    for key in null_keys:
        meta = criteria_meta.get(key, {})
        pending.append(
            {
                "key": key,
                "label": _get_criterion_label(key, criteria_meta),
                "policy_source": meta.get("policy_source", ""),
            }
        )

    focus_key = (state or {}).get("next_target_variable") or (state or {}).get("last_asked_key") or ""
    pending_ask_keys = list((state or {}).get("pending_ask_keys") or [])
    questions_asked = (state or {}).get("questions_asked") or []
    last_question = questions_asked[-1] if questions_asked else ""
    if not last_question and focus_key in checklist:
        last_question = _get_criterion_label(focus_key, criteria_meta)

    conversation_text = _gather_user_conversation_text(chat_instance)
    opening_situation = _opening_situation_text(chat_instance)

    heuristic_updates = _heuristic_checklist_updates(
        user_message,
        checklist,
        criteria_meta,
        focus_key=focus_key,
        pending_ask_keys=pending_ask_keys,
        conversation_text=conversation_text,
    )

    pending_for_prompt = []
    keys_for_context = pending_ask_keys or ([focus_key] if focus_key else [])
    for key in keys_for_context:
        if key not in checklist or checklist.get(key) is not None:
            continue
        meta = criteria_meta.get(key) or {}
        pending_for_prompt.append(
            {
                "key": key,
                "label": _get_criterion_label(key, criteria_meta),
                "policy_source": meta.get("policy_source", ""),
            }
        )

    user_prompt = (
        f"=== PENDING CRITERIA (null / unknown) ===\n"
        f"{json.dumps(pending, indent=2, ensure_ascii=False)}\n\n"
    )
    if pending_for_prompt:
        user_prompt += (
            "=== CRITERIA THE ASSISTANT MOST RECENTLY ASKED ABOUT ===\n"
            f"{json.dumps(pending_for_prompt, indent=2, ensure_ascii=False)}\n\n"
        )
    if opening_situation:
        user_prompt += f"=== OPENING SITUATION (launchpad) ===\n{opening_situation}\n\n"
    if conversation_text:
        user_prompt += f"=== FULL USER CONVERSATION SO FAR ===\n{conversation_text}\n\n"
    if last_question and not pending_for_prompt:
        user_prompt += (
            "=== LAST QUESTION THE USER IS LIKELY ANSWERING ===\n"
            f"{last_question.strip()}\n\n"
        )
    user_prompt += (
        f"=== LATEST USER MESSAGE (prioritize this answer) ===\n{user_message.strip()}\n\n"
        "Return updates JSON for every pending criterion you can resolve from the conversation."
    )

    messages = [
        {"role": "system", "content": EXTRACTOR_SYSTEM_PROMPT},
        {"role": "user", "content": user_prompt},
    ]

    llm_updates = {}
    try:
        raw_content = _call_featherless(messages, max_tokens=512, temperature=0.1)
        parsed = _parse_extractor_response(raw_content)
        for key, value in parsed.items():
            safe_key = _safe_criterion_key(key)
            if safe_key in checklist:
                llm_updates[safe_key] = value
    except (ValueError, RuntimeError):
        logger.exception("Extractor failed for chat %s", chat_instance.id)
        if not heuristic_updates:
            return {"updates": {}, "error": True}

    merged = _merge_extractor_updates(heuristic_updates, llm_updates)
    return {"updates": merged, "error": False}


def _parse_extractor_response(raw_text):
    text = raw_text.strip()
    try:
        data = json.loads(text)
    except json.JSONDecodeError:
        brace_match = re.search(r"(\{.*\})", text, re.DOTALL)
        if not brace_match:
            return {}
        try:
            data = json.loads(brace_match.group(1))
        except json.JSONDecodeError:
            return {}

    updates = data.get("updates") if isinstance(data, dict) else {}
    if not isinstance(updates, dict):
        return {}

    result = {}
    for key, value in updates.items():
        safe_key = _safe_criterion_key(key)
        if not safe_key:
            continue
        sanitized = _sanitize_checklist_value(value)
        if sanitized is not None:
            result[safe_key] = sanitized
    return result


def build_interviewer_user_message(state, *, is_first_turn=False, newly_resolved_keys=None, language_code="en"):
    """Speaker — template-based user-facing copy (no LLM).

    Returns (message, pending_ask_keys) where pending_ask_keys are the criteria
    this turn is asking about (for the extractor on the next user reply).
    """
    language_code = normalize_language_code(language_code)
    checklist = (state or {}).get("checklist") or {}
    criteria_meta = (state or {}).get("criteria_meta") or {}
    null_keys = _null_checklist_keys(checklist)
    askable_keys = _get_askable_null_keys(checklist, criteria_meta)
    newly_resolved_keys = newly_resolved_keys or []
    pending_ask_keys = []

    with override(language_code):
        parts = []
        if is_first_turn:
            parts.append(_("I've reviewed your policy document against what you've shared so far."))
        elif newly_resolved_keys:
            if len(newly_resolved_keys) == 1:
                label = _get_thanks_label(newly_resolved_keys[0], criteria_meta)
                parts.append(_("Thanks — that clarifies %(label)s.") % {"label": label.lower()})
            else:
                parts.append(_("Thanks — that helps clarify several of these points."))

        if not askable_keys:
            return "\n\n".join(parts), pending_ask_keys

        use_single = (not is_first_turn) or len(askable_keys) <= 2

        if use_single:
            target = askable_keys[0]
            pending_ask_keys = [target]
            label = _get_criterion_label(target, criteria_meta).strip().rstrip("?")
            parts.append(_("Can you confirm: %(label)s?") % {"label": label.lower()})
        else:
            pending_ask_keys = askable_keys[:6]
            parts.append(
                _("To check your eligibility based on this document, I need to know a few specific things:")
            )
            for key in pending_ask_keys:
                parts.append(f"* {_get_criterion_label(key, criteria_meta)}")
            if len(askable_keys) > 6:
                parts.append(
                    _("* …and %(count)s other requirement(s).")
                    % {"count": len(askable_keys) - 6}
                )
            parts.append(_("Please provide these details in your next message."))

        return "\n\n".join(parts), pending_ask_keys


def _checklist_has_nulls(eligibility_state):
    checklist = (eligibility_state or {}).get("checklist") or {}
    if not checklist:
        return False
    return any(value is None for value in checklist.values())


def is_interviewer_first_turn(chat_instance):
    """True for the first navigator/interviewer AI turn after has_valid_document."""
    return _count_prior_navigator_ai_messages(chat_instance) == 0


def apply_interviewer_checklist(
    existing_state,
    updated_checklist,
    *,
    user_facing_message=None,
    mark_bulk_completed=False,
    pending_ask_keys=None,
    last_asked_key="",
):
    """Finalize interviewer turn state after extractor + speaker."""
    existing = existing_state if isinstance(existing_state, dict) else {}
    prior = dict(existing.get("checklist") or {})
    sanitized = _sanitize_checklist(updated_checklist) if updated_checklist else {}

    for key in prior:
        if key in sanitized:
            prior[key] = sanitized[key]

    new_state = _build_state_from_checklist(
        existing,
        prior,
        criteria_meta=existing.get("criteria_meta"),
        resolved_sources=existing.get("resolved_sources"),
    )
    new_state["interviewer_bulk_completed"] = (
        existing.get("interviewer_bulk_completed", False) or mark_bulk_completed
    )

    questions = list(existing.get("questions_asked") or [])
    criteria_meta = existing.get("criteria_meta") or {}
    if pending_ask_keys:
        for key in pending_ask_keys:
            label = _get_criterion_label(key, criteria_meta).strip()
            if label and label not in questions:
                questions.append(label)
    elif user_facing_message and "?" in user_facing_message:
        questions.append(user_facing_message.strip())
    new_state["questions_asked"] = questions
    if pending_ask_keys is not None:
        new_state["pending_ask_keys"] = list(pending_ask_keys)
        if pending_ask_keys:
            new_state["next_target_variable"] = pending_ask_keys[0]
    if last_asked_key:
        new_state["last_asked_key"] = last_asked_key
    elif pending_ask_keys and len(pending_ask_keys) == 1:
        new_state["last_asked_key"] = pending_ask_keys[0]

    return _normalize_eligibility_state(new_state)


def _null_checklist_keys(checklist):
    return [key for key, value in (checklist or {}).items() if value is None]


def run_interviewer_loop(chat_instance, user_message="", *, is_first_turn=False, language_code="en"):
    """
    Phase 2 Interviewer — extractor (LLM) + speaker (templates).

    Returns dict: user_facing_message, eligibility_state, error, skip (optional).
    """
    language_code = normalize_language_code(language_code)

    chat_instance.refresh_from_db()
    if not chat_instance.has_valid_document:
        raise ValueError("Interviewer requires chat.has_valid_document=True")

    state = dict(chat_instance.eligibility_state or {})
    if not state.get("planner_completed") or not state.get("checklist"):
        initialize_checklist(chat_instance, language_code=language_code)
        chat_instance.refresh_from_db()
        state = dict(chat_instance.eligibility_state or {})

    newly_resolved = []
    if not is_first_turn and user_message.strip():
        extraction = extract_checklist_updates(
            chat_instance, user_message, state, language_code=language_code
        )
        if not extraction.get("error") and extraction.get("updates"):
            prior = dict(state.get("checklist") or {})
            state = apply_checklist_updates(state, extraction["updates"], source="conversation")
            for key, value in (state.get("checklist") or {}).items():
                if prior.get(key) is None and value is not None:
                    newly_resolved.append(key)

        state["pending_ask_keys"] = []
        state["last_asked_key"] = ""

    state = finalize_eligibility_state(state)

    if not _interview_needs_more_questions(state):
        return {
            "user_facing_message": "",
            "eligibility_state": state,
            "error": False,
            "skip": True,
        }

    user_facing, pending_ask_keys = build_interviewer_user_message(
        state,
        is_first_turn=is_first_turn,
        newly_resolved_keys=newly_resolved,
        language_code=language_code,
    )

    state = apply_interviewer_checklist(
        state,
        state.get("checklist") or {},
        user_facing_message=user_facing,
        mark_bulk_completed=is_first_turn and len(pending_ask_keys) > 2,
        pending_ask_keys=pending_ask_keys,
        last_asked_key=pending_ask_keys[0] if len(pending_ask_keys) == 1 else "",
    )

    return {
        "user_facing_message": user_facing,
        "eligibility_state": state,
        "error": False,
    }


def initialize_checklist(chat_instance, *, language_code="en"):
    """
    Phase 2 Planner — runs once after has_valid_document becomes True.

    Silently extracts qualifying rules from the verified policy, pre-fills known
    answers from the user's opening situation, and saves to chat.eligibility_state.
    No user-facing message is created.
    """
    language_code = normalize_language_code(language_code)

    chat_instance.refresh_from_db()
    if not chat_instance.has_valid_document:
        raise ValueError("Planner requires chat.has_valid_document=True")

    existing_state = chat_instance.eligibility_state or {}
    if existing_state.get("planner_completed") and existing_state.get("checklist"):
        return existing_state

    verified_policy = _build_verified_policy_document(chat_instance)
    first_user_message = chat_instance.messages.filter(role="user").order_by("timestamp").first()
    situation = first_user_message.content.strip() if first_user_message else ""

    if not verified_policy.strip():
        logger.warning("Planner skipped — no extractable policy text for chat %s", chat_instance.id)
        state = _checklist_to_navigator_fields({})
        state["planner_completed"] = True
        state["summary"] = "Policy text unavailable; navigator will rely on conversation."
        chat_instance.eligibility_state = state
        chat_instance.save(update_fields=["eligibility_state"])
        return state

    user_prompt = (
        f"Language directive: criterion labels must be written in "
        f"{language_name_for_prompt(language_code)}.\n\n"
        "=== POLICY DOCUMENT ===\n"
        f"{verified_policy}\n\n"
        "=== USER'S INITIAL SITUATION ===\n"
        f"{situation}\n\n"
        "Return the criteria JSON array with key, label, policy_source, and value for each qualifying rule."
    )

    messages = [
        {"role": "system", "content": planner_system_prompt(language_code)},
        {"role": "user", "content": user_prompt},
    ]

    try:
        raw_content = _call_featherless(messages, max_tokens=1536, temperature=0.1)
        checklist, criteria_meta = _parse_planner_response(raw_content)
        if not checklist:
            logger.warning("Planner returned empty checklist for chat %s", chat_instance.id)

        resolved_sources = {}
        for key, value in checklist.items():
            if value is not None:
                resolved_sources[key] = "opening_situation"

        state = _checklist_to_navigator_fields(checklist, criteria_meta, resolved_sources)
        chat_instance.eligibility_state = state
        chat_instance.save(update_fields=["eligibility_state"])
        return state
    except (ValueError, RuntimeError):
        logger.exception("Planner failed for chat %s", chat_instance.id)
        raise


def _build_verified_policy_document(chat_instance):
    """Extract full verified policy text — prefers cached copy after Gatekeeper approval."""
    chat_instance.refresh_from_db()
    if chat_instance.cached_policy_text:
        return chat_instance.cached_policy_text

    first_user_message = chat_instance.messages.filter(role="user").order_by("timestamp").first()
    if not first_user_message:
        return ""

    return _build_verified_policy_document_from_message(first_user_message)


def _normalize_eligibility_state(state):
    """Ensure eligibility_state has required keys for downstream UI and memory."""
    if not isinstance(state, dict):
        return {}

    normalized = dict(state)
    normalized.setdefault("planner_completed", False)
    normalized.setdefault("checklist", {})
    normalized.setdefault("criteria_meta", {})
    normalized.setdefault("resolved_sources", {})
    normalized.setdefault("interviewer_bulk_completed", False)
    normalized.setdefault("navigation_complete", False)
    normalized.setdefault("conclusion_reached", False)
    normalized.setdefault("phase", "clarification")
    normalized.setdefault("bulk_intake_completed", True)
    normalized.setdefault("summary", "")
    normalized.setdefault("candidate_programs", [])
    normalized.setdefault("criteria", [])
    normalized.setdefault("missing_items", [])
    normalized.setdefault("questions_asked", [])
    normalized.setdefault("pending_ask_keys", [])
    normalized.setdefault("last_asked_key", "")
    normalized.setdefault("next_target_variable", "")
    normalized.setdefault("action_blueprint", "")
    normalized.setdefault("reasoning_trace", "")

    if normalized.get("conclusion_reached") or normalized.get("action_blueprint"):
        if normalized.get("action_blueprint") and not normalized.get("conclusion_reached"):
            normalized["conclusion_reached"] = True
        if normalized.get("conclusion_reached"):
            normalized["phase"] = "blueprint"
            if normalized.get("action_blueprint"):
                normalized["navigation_complete"] = True
        normalized["next_target_variable"] = ""
        normalized["missing_items"] = []
    elif normalized.get("bulk_intake_completed") and normalized.get("phase") == "bulk_intake":
        normalized["phase"] = "clarification"
    elif normalized.get("phase") == "initial":
        normalized["phase"] = "clarification"

    if not isinstance(normalized["questions_asked"], list):
        normalized["questions_asked"] = []
    if not isinstance(normalized["candidate_programs"], list):
        normalized["candidate_programs"] = []
    if not isinstance(normalized["criteria"], list):
        normalized["criteria"] = []
    if not isinstance(normalized["missing_items"], list):
        normalized["missing_items"] = []

    return normalized


_FORBIDDEN_RESPONSE_LINE_PATTERNS = [
    re.compile(r"\beligibility_state\b", re.I),
    re.compile(r"\bmissing_items\b", re.I),
    re.compile(r"\bquestions_asked\b", re.I),
    re.compile(r"\bnext_target_variable\b", re.I),
    re.compile(r"\bbulk_intake\b", re.I),
    re.compile(r"\bknown variables?\b", re.I),
    re.compile(r"\bunknown variables?\b", re.I),
    re.compile(r"\bupdating my (?:checklist|records|state)\b", re.I),
    re.compile(r"\bmy internal (?:checklist|state)\b", re.I),
    re.compile(r"\bI(?:'m| am) (?:tracking|updating|maintaining)\b", re.I),
    re.compile(r"^#{1,3}\s*information I need from you", re.I),
]

_FORBIDDEN_RESPONSE_PHRASES = [
    "eligibility_state",
    "missing_items",
    "questions_asked",
    "next_target_variable",
    "bulk_intake_completed",
    "candidate_programs",
    "policy_source",
]


def _sanitize_user_response(text):
    """Strip internal state narration that leaked into the user-facing response."""
    if not text:
        return text

    lines = text.splitlines()
    cleaned_lines = []
    for line in lines:
        if any(pattern.search(line) for pattern in _FORBIDDEN_RESPONSE_LINE_PATTERNS):
            logger.warning("Removed state-leak line from AI response: %s", line[:120])
            continue
        cleaned_lines.append(line)

    result = "\n".join(cleaned_lines).strip()
    for phrase in _FORBIDDEN_RESPONSE_PHRASES:
        if phrase.lower() in result.lower():
            result = re.sub(re.escape(phrase), "", result, flags=re.I).strip()

    result = re.sub(r"\n{3,}", "\n\n", result).strip()
    return result or text.strip()


def _parse_title_response(raw_text):
    text = raw_text.strip()
    try:
        data = json.loads(text)
        title = data.get("title", "")
    except json.JSONDecodeError:
        brace_match = re.search(r"(\{.*\})", text, re.DOTALL)
        if brace_match:
            try:
                data = json.loads(brace_match.group(1))
                title = data.get("title", "")
            except json.JSONDecodeError:
                title = text.strip('"').strip()
        else:
            title = text.strip('"').strip()

    title = re.sub(r"\s+", " ", str(title)).strip().strip('"').strip("'")
    if not title:
        return None
    if len(title) > MAX_TITLE_LENGTH:
        title = title[: MAX_TITLE_LENGTH - 1].rstrip() + "…"
    return title


def _call_featherless(messages, *, max_tokens=2048, temperature=0.7):
    api_key = os.getenv("FEATHERLESS_API_KEY")
    if not api_key:
        raise ValueError("FEATHERLESS_API_KEY is not configured.")

    model = os.getenv("FEATHERLESS_MODEL", DEFAULT_MODEL)
    headers = {
        "Authorization": f"Bearer {api_key}",
        "Content-Type": "application/json",
        "HTTP-Referer": os.getenv("FEATHERLESS_HTTP_REFERER", "http://localhost:8000"),
        "X-Title": os.getenv("FEATHERLESS_APP_TITLE", "Konya"),
    }
    payload = {
        "model": model,
        "messages": messages,
        "temperature": temperature,
        "max_tokens": max_tokens,
    }

    try:
        response = requests.post(
            FEATHERLESS_API_URL,
            headers=headers,
            json=payload,
            timeout=int(os.getenv("FEATHERLESS_TIMEOUT", "120")),
        )
        response.raise_for_status()
        data = response.json()
        return data["choices"][0]["message"]["content"].strip()
    except requests.RequestException as exc:
        logger.exception("Featherless AI request failed")
        detail = ""
        if exc.response is not None:
            try:
                detail = exc.response.json().get("error", {}).get("message", "")
            except ValueError:
                detail = exc.response.text[:200]
        raise RuntimeError(detail or str(exc)) from exc


def generate_chat_title(first_message_content, language_code="en"):
    """Generate a short 3–5 word title from the user's opening message."""
    language_code = normalize_language_code(language_code)
    language_name = language_name_for_prompt(language_code)

    messages = [
        {"role": "system", "content": TITLE_SYSTEM_PROMPT},
        {
            "role": "user",
            "content": (
                f"Language: {language_name}\n\n"
                f"Opening message:\n{first_message_content[:2000]}"
            ),
        },
    ]

    raw_content = _call_featherless(messages, max_tokens=40, temperature=0.3)
    return _parse_title_response(raw_content)


GUIDE_GENERATOR_SYSTEM_PROMPT = """You are a Navigation Assistant. The user's eligibility interview is complete. \
Review their final eligibility_state against the provided Policy Document. You must generate a final, \
comprehensive Action Blueprint formatted in strict Markdown.

Required structure (use these exact headings):

### Conclusion
A clear statement on whether they may appear eligible based on the document. Use hedged language only \
(e.g. "you may appear eligible", "you may not appear eligible") — never state guaranteed eligibility.

### What You Need to Prepare
List all physical logistics mentioned in the document (e.g., National ID, bank statements, printed forms).

### Next Steps
Clear instructions on where they need to go or who they need to contact to submit their actual application.

End the message with a brief reminder that this is a guide, not an official application, and they must submit \
their paperwork to the relevant authority.

Respond with Markdown only — no JSON, no code fences wrapping the entire response."""


FOLLOW_UP_SYSTEM_PROMPT = """You are Konya providing post-navigation follow-up support.

The user has already received their Action Blueprint. They may ask clarifying questions about the policy \
or next steps.

STRICT RULES:
1. Answer ONLY using the verified policy document provided in this conversation. Do not rely on outside \
knowledge or assumptions beyond what the document states.
2. You may answer simple, obvious questions only when there is zero risk of misleading the user.
3. If the document does not clearly support an answer, politely explain that you cannot answer safely from \
this document alone. Suggest where they might find the correct answer — for example an official website, \
program office, helpline, or contact details mentioned in the document.
4. Do not restart eligibility interviews, change prior conclusions, or offer to submit applications on \
the user's behalf.
5. Keep answers concise, practical, and grounded in the document text.

Respond with Markdown only — no JSON, no code fences wrapping the entire response."""


def is_navigation_complete(chat_instance):
    """True when the Action Blueprint has been issued and navigation is finished."""
    state = chat_instance.eligibility_state or {}
    return bool(
        state.get("navigation_complete")
        or (state.get("conclusion_reached") and state.get("action_blueprint"))
    )


def _apply_blueprint_to_state(existing_state, blueprint_markdown, *, reasoning_trace=""):
    """Mark eligibility_state as complete with the generated Action Blueprint."""
    state = dict(existing_state or {})
    blueprint = blueprint_markdown.strip()
    state["conclusion_reached"] = True
    state["navigation_complete"] = True
    state["phase"] = "blueprint"
    state["action_blueprint"] = blueprint
    state["reasoning_trace"] = (reasoning_trace or "").strip()
    state["summary"] = "Navigation complete — Action Blueprint issued."
    state["missing_items"] = []
    state["next_target_variable"] = ""
    return _normalize_eligibility_state(state)


def is_blueprint_message(message, chat):
    """True when an AI message is the issued Action Blueprint."""
    if message.role != "ai" or message.is_error or is_gatekeeper_message(message):
        return False
    if not is_navigation_complete(chat):
        return False
    content = (message.content or "").strip()
    if not content:
        return False
    stored_blueprint = ((chat.eligibility_state or {}).get("action_blueprint") or "").strip()
    if stored_blueprint and content == stored_blueprint:
        return True
    return "### Conclusion" in content and "### What You Need to Prepare" in content


def _checklist_assessment_label(value, language_code="en"):
    language_code = normalize_language_code(language_code)
    with override(language_code):
        if value is True:
            return _("appears met")
        if value is False:
            return _("appears not met")
        return _("unclear")


def _resolved_source_label(source_key, language_code="en"):
    language_code = normalize_language_code(language_code)
    with override(language_code):
        labels = {
            "opening_situation": _("from your opening description"),
            "conversation": _("from your answers in this chat"),
        }
        return labels.get(source_key, _("from information you provided"))


def build_reasoning_trace_fallback(chat_instance, *, language_code="en"):
    """Deterministic reasoning trace when the LLM call is unavailable."""
    language_code = normalize_language_code(language_code)
    state = chat_instance.eligibility_state or {}
    checklist = state.get("checklist") or {}
    criteria_meta = state.get("criteria_meta") or {}
    resolved_sources = state.get("resolved_sources") or {}

    if not checklist:
        with override(language_code):
            return _("No eligibility criteria were recorded for this navigation.")

    first_user_message = chat_instance.messages.filter(role="user").order_by("timestamp").first()
    opening_situation = first_user_message.content.strip() if first_user_message else ""

    with override(language_code):
        lines = [
            _(
                "This summary shows how Konya matched what you shared against the policy "
                "document you provided."
            ),
            "",
        ]

        for key, value in checklist.items():
            label = _get_criterion_label(key, criteria_meta)
            assessment = _checklist_assessment_label(value, language_code)
            policy_source = (criteria_meta.get(key) or {}).get("policy_source", "").strip()
            source_note = _resolved_source_label(resolved_sources.get(key, ""), language_code)

            item = f"- **{label}** — {assessment} ({source_note})."
            if policy_source:
                item += f' {_("- Policy:")} "{policy_source}"'
            lines.append(item)

        lines.extend(
            [
                "",
                _(
                    "This is based on the policy document you provided. Confirm the outcome "
                    "with the official program office before you act."
                ),
            ]
        )

        if opening_situation:
            lines.insert(
                2,
                _("**Your situation:** %(situation)s")
                % {"situation": opening_situation[:500] + ("…" if len(opening_situation) > 500 else "")},
            )
            lines.insert(3, "")

    return "\n".join(lines)


REASONING_GENERATOR_SYSTEM_PROMPT = """You explain how an eligibility conclusion was reached for a policy navigation tool.

You receive the resolved eligibility checklist, policy source quotes, the user's opening situation, \
conversation history, and the Action Blueprint conclusion.

Write a plain-language Markdown explanation for the user. Structure:

1. One short opening paragraph summarizing the overall logic (use hedged language: "may appear eligible").
2. A bullet list with one item per criterion:
   - **{requirement}** — {met / not met}: what the user shared, what the policy says (quote when available), \
and how that led to the assessment.
3. A closing sentence reminding them to confirm with the official program office.

Rules:
- Never guarantee eligibility or outcomes.
- Only use facts from the provided checklist, conversation, and policy text.
- Do not invent user details or policy rules.
- Keep it concise and practical.

Respond with Markdown only — no JSON, no code fences wrapping the entire response."""


def generate_reasoning_trace(chat_instance, *, language_code="en", blueprint_markdown=""):
    """
    Generate a plain-language trace explaining how the conclusion was reached.

    Returns markdown string (may use deterministic fallback on failure).
    """
    language_code = normalize_language_code(language_code)
    language_name = language_name_for_prompt(language_code)

    state = chat_instance.eligibility_state or {}
    existing = (state.get("reasoning_trace") or "").strip()
    if existing:
        return existing

    checklist = state.get("checklist") or {}
    if not checklist:
        return build_reasoning_trace_fallback(chat_instance, language_code=language_code)

    verified_policy = _build_verified_policy_document(chat_instance)
    criteria_meta = state.get("criteria_meta") or {}
    resolved_sources = state.get("resolved_sources") or {}
    first_user_message = chat_instance.messages.filter(role="user").order_by("timestamp").first()
    opening_situation = first_user_message.content.strip() if first_user_message else ""

    criteria_payload = []
    for key, value in checklist.items():
        meta = criteria_meta.get(key) or {}
        criteria_payload.append(
            {
                "key": key,
                "label": _get_criterion_label(key, criteria_meta),
                "assessment": value,
                "policy_source": meta.get("policy_source", ""),
                "resolved_from": resolved_sources.get(key, ""),
            }
        )

    conversation_lines = []
    for message in chat_instance.messages.order_by("timestamp"):
        if message.role == "ai" and (_is_gatekeeper_message(message) or message.is_error):
            continue
        speaker = "User" if message.role == "user" else "Assistant"
        conversation_lines.append(f"{speaker}: {message.content.strip()}")

    user_prompt = (
        f"Language directive: write the explanation entirely in {language_name}.\n\n"
        f"=== RESOLVED CRITERIA ===\n"
        f"{json.dumps(criteria_payload, indent=2, ensure_ascii=False)}\n\n"
        f"=== OPENING SITUATION ===\n{opening_situation}\n\n"
    )

    if blueprint_markdown.strip():
        user_prompt += f"=== ACTION BLUEPRINT (CONCLUSION) ===\n{blueprint_markdown.strip()}\n\n"

    if verified_policy.strip():
        user_prompt += f"=== VERIFIED POLICY DOCUMENT ===\n{verified_policy}\n\n"

    if conversation_lines:
        user_prompt += (
            "=== CONVERSATION HISTORY ===\n"
            + "\n\n".join(conversation_lines[-20:])
            + "\n\n"
        )

    user_prompt += "Explain how we reached the conclusion."

    messages = [
        {"role": "system", "content": REASONING_GENERATOR_SYSTEM_PROMPT},
        {"role": "user", "content": user_prompt},
    ]

    try:
        raw_content = _call_featherless(
            messages,
            max_tokens=int(os.getenv("FEATHERLESS_REASONING_MAX_TOKENS", "1024")),
            temperature=0.2,
        )
        trace = raw_content.strip()
        if trace:
            return trace
    except (ValueError, RuntimeError):
        logger.exception("Reasoning trace generation failed for chat %s", chat_instance.id)

    return build_reasoning_trace_fallback(chat_instance, language_code=language_code)


def get_or_generate_reasoning_trace(chat_instance, *, language_code="en"):
    """Return stored reasoning trace, generating and persisting it when missing."""
    language_code = normalize_language_code(language_code)
    chat_instance.refresh_from_db()
    state = chat_instance.eligibility_state or {}

    if not is_navigation_complete(chat_instance):
        with override(language_code):
            return "", _("Navigation is not complete yet.")

    existing = (state.get("reasoning_trace") or "").strip()
    if existing:
        return existing, ""

    blueprint = (state.get("action_blueprint") or "").strip()
    trace = generate_reasoning_trace(
        chat_instance,
        language_code=language_code,
        blueprint_markdown=blueprint,
    )
    updated_state = dict(state)
    updated_state["reasoning_trace"] = trace
    chat_instance.eligibility_state = _normalize_eligibility_state(updated_state)
    chat_instance.save(update_fields=["eligibility_state"])
    return trace, ""


def generate_action_blueprint(chat_instance, language_code="en"):
    """
    Phase 2 Guide Generator — runs when all checklist criteria are resolved.

    Returns dict: response (markdown str), eligibility_state (dict), error (bool).
    """
    language_code = normalize_language_code(language_code)
    language_name = language_name_for_prompt(language_code)

    chat_instance.refresh_from_db()
    if not chat_instance.has_valid_document:
        raise ValueError("Guide Generator requires chat.has_valid_document=True")

    state = chat_instance.eligibility_state or {}
    if _checklist_has_nulls(state):
        logger.warning(
            "Guide Generator skipped for chat %s — unresolved checklist keys: %s",
            chat_instance.id,
            _null_checklist_keys(state.get("checklist") or {}),
        )
        return {
            "response": get_ai_error_message(language_code),
            "eligibility_state": state,
            "error": True,
        }

    if is_navigation_complete(chat_instance):
        return {
            "response": state.get("action_blueprint", ""),
            "eligibility_state": state,
            "error": False,
        }

    verified_policy = _build_verified_policy_document(chat_instance)
    checklist = state.get("checklist") or {}
    first_user_message = chat_instance.messages.filter(role="user").order_by("timestamp").first()
    opening_situation = first_user_message.content.strip() if first_user_message else ""

    conversation_lines = []
    for message in chat_instance.messages.order_by("timestamp"):
        if message.role == "ai" and (_is_gatekeeper_message(message) or message.is_error):
            continue
        speaker = "User" if message.role == "user" else "Assistant"
        conversation_lines.append(f"{speaker}: {message.content.strip()}")

    user_prompt = (
        f"Language directive: write the Action Blueprint entirely in {language_name}.\n\n"
        f"=== FINAL ELIGIBILITY STATE (checklist) ===\n"
        f"{json.dumps(checklist, indent=2, ensure_ascii=False)}\n\n"
        f"=== OPENING SITUATION ===\n{opening_situation}\n\n"
    )

    if verified_policy.strip():
        user_prompt += f"=== VERIFIED POLICY DOCUMENT ===\n{verified_policy}\n\n"

    if conversation_lines:
        user_prompt += (
            "=== CONVERSATION HISTORY ===\n"
            + "\n\n".join(conversation_lines)
            + "\n\n"
        )

    user_prompt += "Generate the Action Blueprint now."

    messages = [
        {"role": "system", "content": GUIDE_GENERATOR_SYSTEM_PROMPT},
        {"role": "user", "content": user_prompt},
    ]

    try:
        raw_content = _call_featherless(
            messages,
            max_tokens=int(os.getenv("FEATHERLESS_MAX_TOKENS", "2048")),
            temperature=float(os.getenv("FEATHERLESS_TEMPERATURE", "0.3")),
        )
        blueprint = raw_content.strip()
        if not blueprint:
            raise ValueError("Empty blueprint returned")

        reasoning_trace = generate_reasoning_trace(
            chat_instance,
            language_code=language_code,
            blueprint_markdown=blueprint,
        )
        updated_state = _apply_blueprint_to_state(
            state,
            blueprint,
            reasoning_trace=reasoning_trace,
        )
        return {
            "response": blueprint,
            "eligibility_state": updated_state,
            "error": False,
        }
    except (ValueError, RuntimeError):
        logger.exception("Guide Generator failed for chat %s", chat_instance.id)
        return {
            "response": get_ai_error_message(language_code),
            "eligibility_state": state,
            "error": True,
        }


def generate_follow_up_response(chat_instance, user_message, language_code="en"):
    """
    Post-blueprint Q&A — answers strictly grounded in the verified policy document.
    """
    language_code = normalize_language_code(language_code)
    language_name = language_name_for_prompt(language_code)

    chat_instance.refresh_from_db()
    if not is_navigation_complete(chat_instance):
        raise ValueError("Follow-up requires navigation_complete=True")

    verified_policy = _build_verified_policy_document(chat_instance)
    if not verified_policy.strip():
        with override(language_code):
            return {
                "response": _(
                    "I don't have the policy document available to answer that safely. "
                    "Please check the official program website or contact the relevant office directly."
                ),
                "eligibility_state": chat_instance.eligibility_state or {},
                "error": False,
                "navigation_complete": True,
            }

    state = chat_instance.eligibility_state or {}
    checklist = state.get("checklist") or {}
    blueprint = state.get("action_blueprint", "")

    user_prompt = (
        f"Language directive: write your reply entirely in {language_name}.\n\n"
        f"=== VERIFIED POLICY DOCUMENT ===\n{verified_policy}\n\n"
        f"=== RESOLVED ELIGIBILITY CHECKLIST ===\n"
        f"{json.dumps(checklist, indent=2, ensure_ascii=False)}\n\n"
    )

    if blueprint.strip():
        user_prompt += f"=== ACTION BLUEPRINT ALREADY ISSUED ===\n{blueprint}\n\n"

    user_prompt += f"=== USER FOLLOW-UP QUESTION ===\n{user_message.strip()}"

    messages = [
        {"role": "system", "content": FOLLOW_UP_SYSTEM_PROMPT},
        {"role": "user", "content": user_prompt},
    ]

    try:
        raw_content = _call_featherless(
            messages,
            max_tokens=int(os.getenv("FEATHERLESS_MAX_TOKENS", "1024")),
            temperature=0.2,
        )
        response = _sanitize_user_response(raw_content)
        if not response:
            with override(language_code):
                response = _(
                    "I can't answer that safely from the policy document alone. "
                    "Please check the official program website or contact the relevant office for confirmation."
                )
        return {
            "response": response,
            "eligibility_state": state,
            "error": False,
            "navigation_complete": True,
        }
    except (ValueError, RuntimeError):
        logger.exception("Follow-up response failed for chat %s", chat_instance.id)
        return {
            "response": get_ai_error_message(language_code),
            "eligibility_state": state,
            "error": True,
            "navigation_complete": True,
        }


def generate_chat_ai_response(
    chat_instance,
    *,
    user_message="",
    is_first_turn=None,
    language_code="en",
):
    """
    Orchestrate Phase 2 multi-agent flow.

    Interviewer runs while checklist nulls remain; Guide Generator issues the Action Blueprint
    once all criteria are resolved. After the blueprint, follow-up Q&A is document-grounded only.
    """
    chat_instance.refresh_from_db()
    fallback_state = chat_instance.eligibility_state or {}

    if is_navigation_complete(chat_instance):
        if not user_message.strip():
            return {
                "response": "",
                "eligibility_state": fallback_state,
                "error": False,
                "navigation_complete": True,
                "skip": True,
            }
        return generate_follow_up_response(
            chat_instance,
            user_message,
            language_code=language_code,
        )

    if chat_instance.has_valid_document and not fallback_state.get("planner_completed"):
        try:
            initialize_checklist(chat_instance, language_code=language_code)
            chat_instance.refresh_from_db()
            fallback_state = chat_instance.eligibility_state or {}
        except (ValueError, RuntimeError):
            logger.exception("Planner failed for chat %s", chat_instance.id)

    if is_first_turn is None:
        is_first_turn = is_interviewer_first_turn(chat_instance)

    if chat_instance.has_valid_document and not (fallback_state.get("checklist") or {}):
        return {
            "response": build_empty_planner_message(language_code),
            "eligibility_state": fallback_state,
            "error": False,
            "navigation_complete": False,
        }

    if chat_instance.has_valid_document and _checklist_has_nulls(fallback_state):
        interviewer_result = run_interviewer_loop(
            chat_instance,
            user_message=user_message,
            is_first_turn=is_first_turn,
            language_code=language_code,
        )

        if interviewer_result.get("error"):
            return {
                "response": get_ai_error_message(language_code),
                "eligibility_state": fallback_state,
                "error": True,
            }

        if interviewer_result.get("skip"):
            fallback_state = interviewer_result.get("eligibility_state") or fallback_state
            chat_instance.eligibility_state = fallback_state
            chat_instance.save(update_fields=["eligibility_state"])
            chat_instance.refresh_from_db()
            fallback_state = chat_instance.eligibility_state or fallback_state
        else:
            new_state = interviewer_result["eligibility_state"]
            chat_instance.eligibility_state = new_state
            chat_instance.save(update_fields=["eligibility_state"])
            chat_instance.refresh_from_db()

            if _checklist_has_nulls(new_state):
                return {
                    "response": interviewer_result["user_facing_message"],
                    "eligibility_state": new_state,
                    "error": False,
                    "navigation_complete": False,
                }

            fallback_state = new_state

    if not _checklist_has_nulls(fallback_state) and not is_ready_for_blueprint(fallback_state):
        checklist = fallback_state.get("checklist") or {}
        resolved_sources = dict(fallback_state.get("resolved_sources") or {})
        changed = False
        for key, value in checklist.items():
            if value is not None and key not in resolved_sources:
                resolved_sources[key] = "opening_situation"
                changed = True
        if changed:
            fallback_state = _build_state_from_checklist(
                fallback_state,
                checklist,
                criteria_meta=fallback_state.get("criteria_meta"),
                resolved_sources=resolved_sources,
            )
            chat_instance.eligibility_state = fallback_state
            chat_instance.save(update_fields=["eligibility_state"])

    if is_ready_for_blueprint(fallback_state):
        chat_instance.eligibility_state = fallback_state
        chat_instance.save(update_fields=["eligibility_state"])
        chat_instance.refresh_from_db()
        blueprint_result = generate_action_blueprint(chat_instance, language_code=language_code)
        if not blueprint_result.get("error"):
            chat_instance.eligibility_state = blueprint_result["eligibility_state"]
            chat_instance.save(update_fields=["eligibility_state"])
            return {
                "response": blueprint_result["response"],
                "eligibility_state": blueprint_result["eligibility_state"],
                "error": False,
                "navigation_complete": True,
            }

        return {
            "response": blueprint_result.get("response", get_ai_error_message(language_code)),
            "eligibility_state": fallback_state,
            "error": True,
            "navigation_complete": False,
        }

    return {
        "response": get_ai_error_message(language_code),
        "eligibility_state": fallback_state,
        "error": True,
        "navigation_complete": False,
    }
